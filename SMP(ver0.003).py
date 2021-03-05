from openpyxl import load_workbook
from docx import Document

# *****************МЕНЮ*************************************
file_name = 'saratov.xlsm'  # EXEL Файл где хронится Форма 2
sheet = 'Unit'  # ЛИСТ где СМП
nz = [22, 26, 27]
first_year = '2011'
last_year = '2020'

district_text = 'ПФО'
region_text = 'Саратовской области'
# *****************СЛОВАРЬ*************************************

nz_dict = {1: "Брюшной тиф",
           2: "Паратифы А, В, С и неуточненный",
           3: "Бактерионосители брюшного тифа,паратифов",
           4: "Холера",
           5: "Вибриононосители холеры",
           6: "Другие сальмонеллезные инфекции",
           7: "Сальмонеллы группы B",
           8: "Сальмонеллы группы C",
           9: "Сальмонеллы группы Д",
           10: "Бактериальная дизентерия (шигеллез)",
           11: "из нее бактериологически подтвержденная",
           12: "из нее вызванная: шигеллами Зонне",
           13: "шигеллами Флекснера",
           14: "Бактерионосители дизентерии",
           15: "Другиеострыекишечныеинфекции,вызванныеустановленнымибактериальными,вируснымивозбудителями,атакжепищевыетоксикоинфекцииустановленнойэтиологии",
           16: "изних:вызванныеустановленнымибактериальнымивозбудителями",
           17: "изних:кишечнымипалочками(эшерихиями)",
           18: "кампилобактериями",
           19: "иерсиниями энтероколитика",
           20: "вызванныевирусами",
           21: "изних:ротавирусами",
           22: "Острые кишечные инфекции (ОКИ) вызванные вирусом Норволк",
           23: "Острыекишечныеинфекции,вызванныенеустановленнымиинфекционнымивозбудителями,пищевыетоксикоинфекциинеустановленнойэтиологии",
           24: "Острыйпаралитическийполиомиелит,включаяассоциированныйсвакциной",
           25: "Острыевялыепараличи",
           26: "Энтеровирусныеинфекции",
           27: "изнихэнтеровирусныйменингит",
           28: "Острыегепатиты-всего",
           29: "изних:острыйгепатитА",
           30: "острыйгепатитВ",
           31: "острыйгепатитС",
           32: "острыйгепатитЕ",
           33: "Хроническиевирусныегепатиты(впервыеустановленные)–всего",
           34: "изних:хроническийвирусныйгепатитВ",
           35: "хроническийвирусныйгепатитС",
           36: "***НосительствовозбудителявирусногогепатитаВ",
           37: "Дифтерия",
           38: "Бактерионосителитоксигенныхштаммовдифтерии",
           39: "Коклюш",
           40: "изнегококлюш,вызванныйBordetellaparapertussis",
           41: "Стрептококковаяинфекция(впервыевыявленная)",
           42: "Скарлатина",
           43: "стрептококковаясептицемия",
           44: "Ветрянаяоспа",
           45: "Опоясывающийлишай",
           46: "Корь",
           47: "Краснуха",
           48: "Синдромврожденнойкраснухи(СВК)",
           49: "Паротитэпидемический",
           50: "Генерализованныеформыменингококковойинфекции",
           51: "***Менингококковаяинфекция",
           52: "Гемофильнаяинфекция",
           53: "Столбняк",
           54: "Туляремия",
           55: "Сибирскаяязва",
           56: "Бруцеллез,впервыевыявленный",
           57: "Вирусныелихорадки,передаваемыечленистоногимиивирусныегеморрагическиелихорадки",
           58: "изних:лихорадкаЗападногоНила",
           59: "Крымскаягеморрагическаялихорадка",
           60: "геморрагическаялихорадкаспочечнымсиндромом",
           61: "Омскаягеморрагическаялихорадка",
           62: "лихорадкаденге",
           63: "Клещевой вирусный энцефалит",
           64: "Клещевойборрелиоз(болезньЛайма)",
           65: "Псевдотуберкулез",
           66: "Лептоспироз",
           67: "Бешенство",
           68: "Укусы,ослюнения,оцарапыванияживотными",
           69: "изнихдикимиживотными",
           70: "укусы,нанесенныесобаками",
           71: "Укусыклещами",
           72: "Орнитоз(пситтакоз)",
           73: "Риккетсиозы",
           74: "изних:эпидемическийсыпнойтиф",
           75: "болезньБрилля",
           76: "лихорадкаКу",
           77: "сибирскийклещевойтиф",
           78: "астраханскаяпятнистаялихорадка",
           79: "риккетсиоз,вызываемыйAnaplasmaphagocytophilum",
           80: "риккетсиоз,вызываемыйEhrlichiachaffeensisиEhrlichiamuris",
           81: "Педикулез",
           82: "Листериоз",
           83: "Легионеллез",
           84: "Инфекционныймононуклеоз",
           85: "Туберкулез(впервыевыявленный)активныеформы",
           86: "изнеготуберкулезоргановдыхания",
           87: "изнегобациллярныеформы",
           88: "Сифилис(впервыевыявленный)-всеформы",
           89: "Гонококковаяинфекция",
           90: "Болезнь, вызванная вирусом  иммунодефицита человека (ВИЧ): ",
           91: "Острыеинфекцииверхнихдыхательныхпутеймножественнойинеуточненнойлокализации",
           92: "Грипп",
           93: "Пневмония(внебольничная)",
           94: "изнее:вирусная",
           95: "бактериальная",
           96: "изнее:вызваннаяпневмококками",
           97: "Mycoplasmapneumoniae",
           98: "пневмония,вызваннаяхламидиями",
           99: "Цитомегаловируснаяболезнь",
           100: "Врожденнаяцитомегаловируснаяинфекция",
           101: "Дерматофития,вызваннаягрибамиродаMicrosporum",
           102: "Чесотка",
           103: "Дерматофития,вызваннаягрибамиродаTrichophyton",
           104: "Поствакцинальныеосложнения",
           105: "Паразитоносительствомалярии"}

wb_val = load_workbook(filename=f'{file_name}', data_only=True, )  # Open file
sheet_val = wb_val[f'{sheet}']  # Open sheet
# *****************Вытаскиваем стрчку со значенимя заболевания**********************************************************

abc = ['c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l']  # Столбци таблицы

rus = 3  # Строка с РФ (в Exel)
dis = 4  # Строка с Округом (в Exel)
reg = 5  # Строка с Регионом (в Exel)

russian = []  # Список показателей заболеваемости от нозологии № 1 (С3:L3) до number_nz
district = []  # Список показателей заболеваемости от нозологии № 1 (С4:L4) до number_nz
region = []  # Список показателей заболеваемости от нозологии № 1 (С5:L5) до number_nz

line_russian = []
line_district = []
line_region = []

while rus < 548 or dis < 549 or reg < 550:
    for i in abc:
        ru = sheet_val[f'{i}{rus}'].value
        russian.append(ru)
        di = sheet_val[f'{i}{dis}'].value
        district.append(di)
        re = sheet_val[f'{i}{reg}'].value
        region.append(re)
    rus += 5
    dis += 5
    reg += 5

for number_nz in nz:
    if int(number_nz) == 1:
        c = 0
        l = 10
    else:
        c = 10 * int(number_nz - 1)  # Индекс значения заболеваемости первого года (first_year)
        l = 10 * int(number_nz)  # Индекс значения заболеваемости последнего года (last_year)\

    string_russian = russian[c:l]  # Строчка со значениями заболеваемости нозологии №(nz) по РФ
    string_district = district[c:l]  # Строчка со значениями заболеваемости нозологии №(nz) по Округу
    string_region = region[c:l]  # Строчка со значениями заболеваемости нозологии №(nz) по Региону

    # line_number_nz = (f'Нозологи № {number_nz}')

    # print(string_russian)
    # print(string_district)
    # print(string_region)

    line_russian.append(string_russian)
    line_district.append(string_district)
    line_region.append(string_region)  # [0: (len(string_region) - 1)])

# **************************************Вписываем_значения_заболеваемости_Рф_Окуруга_Региона***************************************************


new_list_rf = []  # список значений заболеваемости за последний год по РФ
new_list_district = []  # список значений заболеваемости за последний год по Округу
new_list_reg = []  # список значений заболеваемости за последний год по Региону
new_list_smp = []


new_list_full_mr_rf = []
new_list_full_mr_ds = []
new_list_full_mr_rg = []


for i in line_russian:
    last_year_nz = i[len(i) - 1]
    new_list_rf.append(last_year_nz)  # Добавляем значений заболеваемости за последний год по РФ

    # for i_full in i:
    #     new_list_full_mr_rf.append(i_full)

for i in line_district:
    last_year_nz = i[len(i) - 1]
    new_list_district.append(last_year_nz)  # Добавляем значений заболеваемости за последний год по Округу

    # for i_full in i:
    #     new_list_full_mr_ds.append(i_full)

for i in line_region:
    last_year_nz = i[len(i) - 1]
    new_list_reg.append(last_year_nz)  # Добавляем значений заболеваемости за последний год по Региону

    # for i_full in i:
    #     new_list_full_mr_rg.append(i_full)


print(line_russian)
# print(len(new_list_full_mr_rf))
# print(new_list_full_mr_ds)
# print(new_list_full_mr_rg)
# **************************************СЧИТАЕМ___СМП*******************************************************


for i in line_region:
    # print(i)
    i.pop()
    smp_region = i
    smp_list = []

    for element in smp_region:
        if element != 0:
            smp_list.append(element)

    smp_list_copy = smp_list.copy()

    maximum = max(smp_list_copy)
    idx_max = smp_list_copy.index(maximum)
    smp_list_copy.pop(idx_max)

    minimum = min(smp_list_copy)
    idx_min = smp_list_copy.index(minimum)
    smp_list_copy.pop(idx_min)

    smp_value = sum(smp_list_copy) / len(smp_list_copy)
    smp = "%.2f" % smp_value
    new_list_smp.append(smp)
# **********************************************ВЫШЕ\НИЖЕ****************************************************
up_down = {1: 'ниже', 2: 'на уровне', 3: 'выше'}
pretext = {1: 'в', 2: 'на', 3: ' '}
time_percent = {1: 'раза.', 2: '%.', 3: ''}

# **********************************************ТЕКСТОВКА*****************************************************

documents = Document()  # Открываем документ

p = documents.add_paragraph()  # Создаем параграф
p.add_run('МАТЕРИАЛЫ').bold = True  # Записываем в этот пораграф
p.add_run(f' о деятельности Управления Роспотребнадзора по {region_text}').bold = True
p.add_run(f' и ФБУЗ «Центр гигиены и эпидемиологии в {region_text}»').bold = True
p = documents.add_paragraph()  # Создаем параграф
p.add_run(f'Инфекционная и паразитарная заболеваемость в {first_year} - {last_year}гг.').bold = True
p = documents.add_paragraph('')  # Создаем параграф
p.add_run(f'В {last_year} году для некоторых инфекционных болезней отмечается подъем заболеваемости, ')
p.add_run('что может быть вызвано как ухудшением эпидемиологической ситуации, ')
p.add_run('так и характерными проявлениями эпидемического процесса отдельных ')
p.add_run('инфекций или улучшением качества диагностики среди населения. ')
p.add_run('Так, в многолетней динамике отмечается рост заболеваемости: ')
# ***************************************ТАБЛИЦА***************************************
table = documents.add_table(rows=5 * len(nz), cols=12, style='Table Grid')
# ГОД
rows_year_ls = []
rows_year_value = 1  # step 5
cols_year_ls = []
year_ls = []
rows_mr_rf = []
rows_mr_rf_value = 2

year_value = int(first_year)

while len(rows_year_ls) < len(nz):
    rows_year_ls.append(rows_year_value)
    rows_mr_rf.append(rows_mr_rf_value)
    rows_year_value += 5
    rows_mr_rf_value += 5

for i in range(10):
    cols_year_ls.append(i + 1)

while len(year_ls) < len(cols_year_ls):
    year_ls.append(year_value)

    year_value += 1

for i_rows in range(len(nz)):
    for i_cols in range(len(cols_year_ls)):
        cell_yaer = table.cell(rows_year_ls[i_rows], cols_year_ls[i_cols])
        cell_yaer.text = f'{year_ls[i_cols]}'

        cell_mr_rf = table.cell(rows_mr_rf[i_rows], cols_year_ls[i_cols])
        # for i in line_russian:

    # cell_mr_rf.text = f'{line_russian[i_cols]}'

# РФ ОКРУГ РЕГИОН + НОЗОЛОГИЯ
first_cell_rf = 2
first_cell_ds = 3
first_cell_rg = 4
first_cell_nz = 0

nz_ls = []
nz_ls_value = 0
for i in nz:
    nz_ls.append(i)
while first_cell_rf <= len(nz) * 5:
    cell_rf = table.cell(first_cell_rf, 0)
    cell_rf.text = 'Российская Федерация'
    cell_ds = table.cell(first_cell_ds, 0)
    cell_ds.text = f'{district_text}'
    cell_rg = table.cell(first_cell_rg, 0)
    cell_rg.text = f'{region_text}'

    cell_nz = table.cell(first_cell_nz, 0)
    cell_nz.text = f'{nz_dict.get(nz_ls[nz_ls_value])}'

    first_cell_rf += 5
    first_cell_ds += 5
    first_cell_rg += 5
    first_cell_nz += 5
    nz_ls_value += 1

# ******************************************ЗАБОЛЕВАЕМОСТЬ В ТАБЛИЦУ
first_cell_mr_rf = 0
first_cell_mr_ds = 0
first_cell_mr_rg = 0

print(string_russian[0])

p = documents.add_paragraph()  # Создаем параграф

for i in range(len(nz)):
    p = documents.add_paragraph()  # Создаем параграф
    p.add_run(f'{nz_dict.get(nz[i])} – ').bold = True
    p.add_run(
        f'{new_list_reg[i]} на 100 тыс. населения при среднемноголетней заболеваемости {new_list_smp[i]}. Показатель ')
    p.add_run(f'по субъекту в {last_year} году')

    up_down_list = []
    pretext_list = []
    time_percent_list = []

    time_percent_value_list = []  # Если < 1.5 раз то %, если > 1.5 то разы

    percent = (new_list_reg[i] * 100) / new_list_rf[i] - 100  # Узнаем в процентах
    time = 0.0  # узнаем во сколько раз больше
    if new_list_reg[i] > new_list_rf[i]:
        time = new_list_reg[i] / new_list_rf[i]
        up_down_list.append(up_down.get(3))  # если регион больше РФ то  показатель заболеваемости (up_down) = ВЫШЕ
    else:
        time = new_list_rf[i] / new_list_reg[i]
        up_down_list.append(up_down.get(1))  # если РФ больше региона то  показатель заболеваемости (up_down) = НИЖЕ

    if abs(time) > 1.5:
        pretext_list.append(pretext.get(1))  # в
        time_value = "%.2f" % abs(time)
        time_percent_value_list.append(time_value)
        time_percent_list.append(time_percent.get(1))

    if abs(time) <= 1.5:
        if 0 < abs(percent) < 12:  # ПРОЦЕНТЫ на 'уровне'
            up_down_list.clear()
            up_down_list.append(up_down.get(2))
            pretext_list.append(pretext.get(3))
            time_percent_value_list.append('')
            time_percent_list.append(time_percent.get(3))
        else:
            pretext_list.append(pretext.get(2))
            percent_value = "%.2f" % abs(percent)
            time_percent_value_list.append(percent_value)
            time_percent_list.append(time_percent.get(2))

    upd = up_down_list[0]  # выше / на уровне / ниже
    pt_l = pretext_list[0]  # в / на
    ttp_value = time_percent_value_list[0]  # значение
    tp_l = time_percent_list[0]  # в % или раз

    p.add_run(f' {upd} ').bold = True
    p.add_run(f'показателя по Российской Федерации ({new_list_rf[i]} на 100 тыс. населения) ')
    p.add_run(f'{pt_l} {time_percent_value_list[0]} {tp_l}').bold = True
    p.add_run(
        f'  Заболеваемость {nz_dict.get(nz[i])} в  {district_text} в {last_year} составила {new_list_district[i]} на 100 тыс. населения')
    p = documents.add_paragraph()  # Создаем параграф
    p.add_run(
        f'Рис.{i + 1} Заболеваемость {nz_dict.get(nz[i])} в {region_text} в {first_year} - {last_year}гг. (на 100 тыс. населения).')

documents.save(f'xсправкапо{region_text}.docx')  # ПОСЛЕДНЯЯ СТРОЧКА В КОДЕ!!!
