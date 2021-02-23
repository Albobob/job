from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches

# Menu
file_name = 'sprvk.xlsx'  # EXEL Файл где хронится Форма 2
sheet = 'Unit'  # ЛИСТ где СМП
number_nz = '6'
obl = 'Свердовской области'  # Субъект
years_1 = '2010'  # начальный год
years_2 = '2019'  # конецчный год

# ******************************************************************С_Л_О_В_А_Р_Ь(нозологии)****************************************************
nz_dict = {1: 'Брюшной тиф', 2: 'Паратифы А,В,С и неуточненный', 3: 'Бактерионосители брюшного тифа, паратифов',
           4: 'Холера', 5: 'Вибриононосители холеры', 6: 'Другие сальмонеллезные инфекции',
           7: 'из них вызванные: сальмонеллами группы B', 8: 'сальмонеллами группы C', 9: 'сальмонеллами группы Д',
           10: 'Бактериальная дизентерия (шигеллез)', 11: 'из нее бактериологически подтвержденная',
           12: 'из нее вызванная: шигеллами Зонне', 13: 'шигеллами Флекснера', 14: 'Бактерионосители дизентерии',
           15: 'ОКИ установленной этиологии и ПТИ'}

# ******************************************************************с_ч_и_т_а_е_м___С_М_П****************************************************
wb_val = load_workbook(filename=f'{file_name}', data_only=True, )  # Open file

sheet_val = wb_val[f'{sheet}']  # Open sheet
abc = ['c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l']  # Выбираю строку со значениями
new_list = []  # Список со всеми занчениями заболеваемости региона
alb = 5
cell_alb = (alb * int(number_nz) + 1)
while alb < cell_alb:
    for i in abc:
        # print(alb)
        # print(i)
        c5_val = sheet_val[f'{i}{alb}'].value
        new_list.append(c5_val)

    alb += 5
print(f' Все НОЗОЛОГИИ: {new_list}')
# print(len(new_list))    # кол-во ячеек
nz_list = []
cl_5 = 0
lc_5 = 9

while cl_5 <= (int(number_nz) * 10) - 10 and lc_5 <= (int(number_nz) * 10) - 1:
    string = new_list[cl_5:lc_5]
    nz_list.append(string)

    cl_5 += 10
    lc_5 += 10
print(f'Нозологии без последнего года: {nz_list}')
# print(len(nz_list))   # кол-во листов
string_nz = nz_list[int(number_nz) - 1]

print(f'Нозология №{number_nz}: {string_nz}')
string_nz_copy = string_nz.copy()

for i in string_nz_copy:
    if i == 0:
        null = string_nz_copy.index(0)
        string_nz_copy.pop(null)  # Удаляем ноль
maximum = max(string_nz_copy)
minimum = min(string_nz_copy)
indx_max = string_nz_copy.index(maximum)
string_nz_copy.pop(indx_max)  # Удаляем Максимальное число
indx_min = string_nz_copy.index(minimum)
string_nz_copy.pop(indx_min)  # Удаляем минимальное число

smp = sum(string_nz_copy) / len(string_nz_copy)

print(f'{smp} СМП нозологии №{number_nz}')

# ******************************************************************Ш_А_П_К_А****************************************************

documents = Document('text_finish.docx')  # Открываем документ
p = documents.add_paragraph()  # Создаем параграф
p.add_run('МАТЕРИАЛЫ').bold = True  # Записываем в этот пораграф
p.add_run(f' о деятельности Управления Роспотребнадзора по {obl}').bold = True
p.add_run(f' и ФБУЗ «Центр гигиены и эпидемиологии в {obl}»').bold = True
p = documents.add_paragraph()  # Создаем параграф
p.add_run(f'Инфекционная и паразитарная заболеваемость в {years_1} - {years_2}гг.').bold = True
p = documents.add_paragraph('')  # Создаем параграф
p.add_run(f'В {years_2} году для некоторых инфекционных болезней отмечается подъем заболеваемости, ')
p.add_run('что может быть вызвано как ухудшением эпидемиологической ситуации, ')
p.add_run('так и характерными проявлениями эпидемического процесса отдельных ')
p.add_run('инфекций или улучшением качества диагностики среди населения. ')
p.add_run('Так, в многолетней динамике отмечается рост заболеваемости: ')

# ******************************************************************Нозологии****************************************************
nz_one = nz_dict.get(int(number_nz))
nd = '***'  # Надо ВСТАВИТЬ
p = documents.add_paragraph(f'')  # Создаем параграф
p.add_run(f'{nz_one}').bold = True
p.add_run(f' {nd} на 100 тыс. населения при среднемноголетней заболеваемости {smp}')

documents.save('text_finish.docx')  # ПОСЛЕДНЯЯ СТРОЧКА В КОДЕ!!!
