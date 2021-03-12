from openpyxl import load_workbook
from docx import Document

# Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню***Меню
# Работа с файлом EXEL
file_name = 'saratov.xlsm'
sheet = 'Unit'
# Работа с нозологией
nz = [22, 26]
first_year = '2011'
last_year = '2020'
district = 'ПФО'
region = 'Саратовская облость'

# Словари***Словари***Словари***Словари***Словари***Словари***Словари***Словари***Словари***Словари***Словари***Словари

# Словарь_Нозологий
nz_dict = {1: "Брюшной тиф",
           2: "Паратифы А, В, С и неуточненный",
           3: "Бактерионосители брюшного тифа,паратифов",
           4: "Холера",
           5: "Вибриононосители холеры",
           6: "Другие сальмонеллезные инфекции",
           7: "Сальмонеллы группы B",
           8: "Сальмонеллы группы C",
           9: "Сальмонеллы группы Д",
           10: "Бактериальная дизентерия (шигеллез)",
           11: "из нее бактериологически подтвержденная",
           12: "из нее вызванная: шигеллами Зонне",
           13: "шигеллами Флекснера",
           14: "Бактерионосители дизентерии",
           15: "Другиеострыекишечныеинфекции,вызванныеустановленнымибактериальными,вируснымивозбудителями,"
               "атакжепищевыетоксикоинфекцииустановленнойэтиологии",
           16: "изних:вызванныеустановленнымибактериальнымивозбудителями",
           17: "изних:кишечнымипалочками(эшерихиями)",
           18: "кампилобактериями",
           19: "иерсиниями энтероколитика",
           20: "вызванныевирусами",
           21: "изних:ротавирусами",
           22: "Острые кишечные инфекции (ОКИ) вызванные вирусом Норволк",
           23: "Острыекишечныеинфекции,вызванныенеустановленнымиинфекционнымивозбудителями, "
               "пищевыетоксикоинфекциинеустановленнойэтиологии",
           24: "Острыйпаралитическийполиомиелит,включаяассоциированныйсвакциной",
           25: "Острыевялыепараличи",
           26: "Энтеровирусныеинфекции",
           27: "изнихэнтеровирусныйменингит",
           28: "Острыегепатиты-всего",
           29: "изних:острыйгепатитА",
           30: "острыйгепатитВ",
           31: "острыйгепатитС",
           32: "острыйгепатитЕ",
           33: "Хроническиевирусныегепатиты(впервыеустановленные)–всего",
           34: "изних:хроническийвирусныйгепатитВ",
           35: "хроническийвирусныйгепатитС",
           36: "***НосительствовозбудителявирусногогепатитаВ",
           37: "Дифтерия",
           38: "Бактерионосителитоксигенныхштаммовдифтерии",
           39: "Коклюш",
           40: "изнегококлюш,вызванныйBordetellaparapertussis",
           41: "Стрептококковаяинфекция(впервыевыявленная)",
           42: "Скарлатина",
           43: "стрептококковаясептицемия",
           44: "Ветрянаяоспа",
           45: "Опоясывающийлишай",
           46: "Корь",
           47: "Краснуха",
           48: "Синдромврожденнойкраснухи(СВК)",
           49: "Паротитэпидемический",
           50: "Генерализованныеформыменингококковойинфекции",
           51: "***Менингококковаяинфекция",
           52: "Гемофильнаяинфекция",
           53: "Столбняк",
           54: "Туляремия",
           55: "Сибирскаяязва",
           56: "Бруцеллез,впервыевыявленный",
           57: "Вирусныелихорадки,передаваемыечленистоногимиивирусныегеморрагическиелихорадки",
           58: "изних:лихорадкаЗападногоНила",
           59: "Крымскаягеморрагическаялихорадка",
           60: "геморрагическаялихорадкаспочечнымсиндромом",
           61: "Омскаягеморрагическаялихорадка",
           62: "лихорадкаденге",
           63: "Клещевой вирусный энцефалит",
           64: "Клещевойборрелиоз(болезньЛайма)",
           65: "Псевдотуберкулез",
           66: "Лептоспироз",
           67: "Бешенство",
           68: "Укусы,ослюнения,оцарапыванияживотными",
           69: "изнихдикимиживотными",
           70: "укусы,нанесенныесобаками",
           71: "Укусыклещами",
           72: "Орнитоз(пситтакоз)",
           73: "Риккетсиозы",
           74: "изних:эпидемическийсыпнойтиф",
           75: "болезньБрилля",
           76: "лихорадкаКу",
           77: "сибирскийклещевойтиф",
           78: "астраханскаяпятнистаялихорадка",
           79: "риккетсиоз,вызываемыйAnaplasmaphagocytophilum",
           80: "риккетсиоз,вызываемыйEhrlichiachaffeensisиEhrlichiamuris",
           81: "Педикулез",
           82: "Листериоз",
           83: "Легионеллез",
           84: "Инфекционныймононуклеоз",
           85: "Туберкулез(впервыевыявленный)активныеформы",
           86: "изнеготуберкулезоргановдыхания",
           87: "изнегобациллярныеформы",
           88: "Сифилис(впервыевыявленный)-всеформы",
           89: "Гонококковаяинфекция",
           90: "Болезнь, вызванная вирусом  иммунодефицита человека (ВИЧ): ",
           91: "Острыеинфекцииверхнихдыхательныхпутеймножественнойинеуточненнойлокализации",
           92: "Грипп",
           93: "Пневмония(внебольничная)",
           94: "изнее:вирусная",
           95: "бактериальная",
           96: "изнее:вызваннаяпневмококками",
           97: "Mycoplasmapneumoniae",
           98: "пневмония,вызваннаяхламидиями",
           99: "Цитомегаловируснаяболезнь",
           100: "Врожденнаяцитомегаловируснаяинфекция",
           101: "Дерматофития,вызваннаягрибамиродаMicrosporum",
           102: "Чесотка",
           103: "Дерматофития,вызваннаягрибамиродаTrichophyton",
           104: "Поствакцинальныеосложнения",
           105: "Паразитоносительствомалярии"}
# Словарь_Заболеваемости_Нозологий
nz_list_dict_rf = {0: [],
                   1: [],
                   2: [],
                   3: [],
                   4: [],
                   5: [],
                   6: [],
                   7: [],
                   8: [],
                   9: [],
                   10: [],
                   11: [],
                   12: [],
                   13: [],
                   14: [],
                   15: [],
                   16: [],
                   17: [],
                   18: [],
                   19: [],
                   20: [],
                   21: [],
                   22: [],
                   23: [],
                   24: [],
                   25: [],
                   26: [],
                   27: [],
                   28: [],
                   29: [],
                   30: [],
                   31: [],
                   32: [],
                   33: [],
                   34: [],
                   35: [],
                   36: [],
                   37: [],
                   38: [],
                   39: [],
                   40: [],
                   41: [],
                   42: [],
                   43: [],
                   44: [],
                   45: [],
                   46: [],
                   47: [],
                   48: [],
                   49: [],
                   50: [],
                   51: [],
                   52: [],
                   53: [],
                   54: [],
                   55: [],
                   56: [],
                   57: [],
                   58: [],
                   59: [],
                   60: [],
                   61: [],
                   62: [],
                   63: [],
                   64: [],
                   65: [],
                   66: [],
                   67: [],
                   68: [],
                   69: [],
                   70: [],
                   71: [],
                   72: [],
                   73: [],
                   74: [],
                   75: [],
                   76: [],
                   77: [],
                   78: [],
                   79: [],
                   80: [],
                   81: [],
                   82: [],
                   83: [],
                   84: [],
                   85: [],
                   86: [],
                   87: [],
                   88: [],
                   89: [],
                   90: [],
                   91: [],
                   92: [],
                   93: [],
                   94: [],
                   95: [],
                   96: [],
                   97: [],
                   98: [],
                   99: [],
                   100: [],
                   101: [],
                   102: [],
                   103: [],
                   104: [],
                   105: []
                   }
nz_list_dict_ds = {0: [],
                   1: [],
                   2: [],
                   3: [],
                   4: [],
                   5: [],
                   6: [],
                   7: [],
                   8: [],
                   9: [],
                   10: [],
                   11: [],
                   12: [],
                   13: [],
                   14: [],
                   15: [],
                   16: [],
                   17: [],
                   18: [],
                   19: [],
                   20: [],
                   21: [],
                   22: [],
                   23: [],
                   24: [],
                   25: [],
                   26: [],
                   27: [],
                   28: [],
                   29: [],
                   30: [],
                   31: [],
                   32: [],
                   33: [],
                   34: [],
                   35: [],
                   36: [],
                   37: [],
                   38: [],
                   39: [],
                   40: [],
                   41: [],
                   42: [],
                   43: [],
                   44: [],
                   45: [],
                   46: [],
                   47: [],
                   48: [],
                   49: [],
                   50: [],
                   51: [],
                   52: [],
                   53: [],
                   54: [],
                   55: [],
                   56: [],
                   57: [],
                   58: [],
                   59: [],
                   60: [],
                   61: [],
                   62: [],
                   63: [],
                   64: [],
                   65: [],
                   66: [],
                   67: [],
                   68: [],
                   69: [],
                   70: [],
                   71: [],
                   72: [],
                   73: [],
                   74: [],
                   75: [],
                   76: [],
                   77: [],
                   78: [],
                   79: [],
                   80: [],
                   81: [],
                   82: [],
                   83: [],
                   84: [],
                   85: [],
                   86: [],
                   87: [],
                   88: [],
                   89: [],
                   90: [],
                   91: [],
                   92: [],
                   93: [],
                   94: [],
                   95: [],
                   96: [],
                   97: [],
                   98: [],
                   99: [],
                   100: [],
                   101: [],
                   102: [],
                   103: [],
                   104: [],
                   105: []
                   }
nz_list_dict_rg = {0: [],
                   1: [],
                   2: [],
                   3: [],
                   4: [],
                   5: [],
                   6: [],
                   7: [],
                   8: [],
                   9: [],
                   10: [],
                   11: [],
                   12: [],
                   13: [],
                   14: [],
                   15: [],
                   16: [],
                   17: [],
                   18: [],
                   19: [],
                   20: [],
                   21: [],
                   22: [],
                   23: [],
                   24: [],
                   25: [],
                   26: [],
                   27: [],
                   28: [],
                   29: [],
                   30: [],
                   31: [],
                   32: [],
                   33: [],
                   34: [],
                   35: [],
                   36: [],
                   37: [],
                   38: [],
                   39: [],
                   40: [],
                   41: [],
                   42: [],
                   43: [],
                   44: [],
                   45: [],
                   46: [],
                   47: [],
                   48: [],
                   49: [],
                   50: [],
                   51: [],
                   52: [],
                   53: [],
                   54: [],
                   55: [],
                   56: [],
                   57: [],
                   58: [],
                   59: [],
                   60: [],
                   61: [],
                   62: [],
                   63: [],
                   64: [],
                   65: [],
                   66: [],
                   67: [],
                   68: [],
                   69: [],
                   70: [],
                   71: [],
                   72: [],
                   73: [],
                   74: [],
                   75: [],
                   76: [],
                   77: [],
                   78: [],
                   79: [],
                   80: [],
                   81: [],
                   82: [],
                   83: [],
                   84: [],
                   85: [],
                   86: [],
                   87: [],
                   88: [],
                   89: [],
                   90: [],
                   91: [],
                   92: [],
                   93: [],
                   94: [],
                   95: [],
                   96: [],
                   97: [],
                   98: [],
                   99: [],
                   100: [],
                   101: [],
                   102: [],
                   103: [],
                   104: [],
                   105: []
                   }
# Cписки строк заболеваемости
number_string_rus = [
    3, 8, 13, 18, 23, 28, 33, 38, 43, 48, 53, 58, 63, 68, 73, 78, 83, 88, 93, 98, 103, 108, 113, 118, 123, 128, 133,
    138, 143, 148, 153, 158, 163, 168, 173, 178, 183, 188, 193, 198, 203, 208, 213, 218, 223, 228, 233, 238, 243, 248,
    253, 258, 263, 268, 273, 278, 283, 288, 293, 298, 303, 308, 313, 318, 323, 328, 333, 338, 343, 348, 353, 358, 363,
    368, 373, 378, 383, 388, 393, 398, 403, 408, 413, 418, 423, 428, 433, 438, 443, 448, 453, 458, 463, 468, 473, 478,
    483, 488, 493, 498, 503, 508, 513, 518, 523, 528, 533, 538, 543]
number_string_dis = [4, 9, 14, 19, 24, 29, 34, 39, 44, 49, 54, 59, 64, 69, 74, 79, 84, 89, 94, 99, 104, 109, 114, 119,
                     124, 129, 134, 139, 144, 149, 154, 159, 164, 169, 174, 179, 184, 189, 194, 199, 204, 209, 214, 219,
                     224, 229, 234, 239, 244, 249, 254, 259, 264, 269, 274, 279, 284, 289, 294, 299, 304, 309, 314, 319,
                     324, 329, 334, 339, 344, 349, 354, 359, 364, 369, 374, 379, 384, 389, 394, 399, 404, 409, 414, 419,
                     424, 429, 434, 439, 444, 449, 454, 459, 464, 469, 474, 479, 484, 489, 494, 499, 504, 509, 514, 519,
                     524, 529, 534, 539, 544]
number_string_reg = [5, 10, 15, 20, 25, 30, 35, 40, 45, 50, 55, 60, 65, 70, 75, 80, 85, 90, 95, 100, 105, 110, 115, 120,
                     125, 130, 135, 140, 145, 150, 155, 160, 165, 170, 175, 180, 185, 190, 195, 200, 205, 210, 215, 220,
                     225, 230, 235, 240, 245, 250, 255, 260, 265, 270, 275, 280, 285, 290, 295, 300, 305, 310, 315, 320,
                     325, 330, 335, 340, 345, 350, 355, 360, 365, 370, 375, 380, 385, 390, 395, 400, 405, 410, 415, 420,
                     425, 430, 435, 440, 445, 450, 455, 460, 465, 470, 475, 480, 485, 490, 495, 500, 505, 510, 515, 520,
                     525, 530, 535, 540, 545]

# Работа c EXEL
wb = load_workbook(filename=f'{file_name}')
sheet_val = wb[f'{sheet}']

# *****Вытаскиваем стрчку со значенимя заболевания*****
cols_abc = ['c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l']

for i_nz in range(106):
    for i in cols_abc:
        ru = sheet_val[f'{i}{number_string_rus[i_nz]}'].value
        rf_string = nz_list_dict_rf.get(i_nz)
        rf_string.append(ru)

# num = 90
# print(f'Нозологи -{nz_dict.get(num)}'
#       f'Заболеваемость по РФ {nz_list_dict_rf[num - 1]}')

# **********************************************ТЕКСТОВКА*****************************************************
documents = Document()  # Открываем документ

p = documents.add_paragraph()  # Создаем параграф
p.add_run('МАТЕРИАЛЫ').bold = True  # Записываем в этот пораграф
p.add_run(f' о деятельности Управления Роспотребнадзора по {region}').bold = True
p.add_run(f' и ФБУЗ «Центр гигиены и эпидемиологии в {region}»').bold = True
p = documents.add_paragraph()  # Создаем параграф
p.add_run(f'Инфекционная и паразитарная заболеваемость в {first_year} - {last_year}гг.').bold = True
p = documents.add_paragraph('')  # Создаем параграф
p.add_run(f'В {last_year} году для некоторых инфекционных болезней отмечается подъем заболеваемости, ')
p.add_run('что может быть вызвано как ухудшением эпидемиологической ситуации, ')
p.add_run('так и характерными проявлениями эпидемического процесса отдельных ')
p.add_run('инфекций или улучшением качества диагностики среди населения. ')
p.add_run('Так, в многолетней динамике отмечается рост заболеваемости: ')

# ***************************************ТАБЛИЦА***************************************
table = documents.add_table(rows=5 * len(nz), cols=12, style='Table Grid')

rows_sum = 5 * len(nz)

# Cписки строк ГОДА
year_list = []  # Список с годами
first_last_year = int(first_year)

while first_last_year <= int(last_year):
    year_list.append(first_last_year)

    first_last_year += 1

rows_list_year = []  # Список с номерами строчек для записи года
rows_year = 1

while rows_year <= rows_sum:
    rows_list_year.append(rows_year)

    rows_year += 5

cols_list = []  # Список с номерами столбцов для записи года, заболеваемости и тд...
cols = 1

while cols <= 10:
    cols_list.append(cols)

    cols += 1

rows_value = 1
cols_value = 1

# year_xyz = table.cell(, 0)

for i_rows in range(len(nz)):
    for i_cols in range(len(cols_list)):
        cell_yaer = table.cell(rows_list_year[i_rows], cols_list[i_cols])
        cell_yaer.text = f'{year_list[i_cols]}'

rows_list_mr = []  # Список с номерами строчек для записи года
rows_mr = 2

while rows_mr <= rows_sum:
    rows_list_mr.append(rows_mr)


    rows_mr += 5



documents.save(f'demo_demo.docx')  # ПОСЛЕДНЯЯ СТРОЧКА В КОДЕ!!!
